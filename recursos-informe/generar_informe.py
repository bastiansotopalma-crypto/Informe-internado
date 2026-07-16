# -*- coding: utf-8 -*-
"""Informe de Internado en Farmacia Asistencial y APS (CESFAM Villa Nonguen).
Formato basado en la memoria de titulo del estudiante, con las modificaciones
del Instructivo de Internados USS 2026."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
NAVY = RGBColor(0x1f, 0x3b, 0x5c)
BLACK = RGBColor(0, 0, 0)
FIG = "recursos-informe"

doc = Document()

# ------------------------------------------------------------- base styles
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def style_heading(name, align, upper_color=NAVY, space_before=16, space_after=10):
    st = doc.styles[name]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.font.bold = True
    st.font.color.rgb = upper_color
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.alignment = align
    st.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    st.paragraph_format.space_before = Pt(space_before)
    st.paragraph_format.space_after = Pt(space_after)
    st.paragraph_format.keep_with_next = True

style_heading("Heading 1", WD_ALIGN_PARAGRAPH.CENTER)
style_heading("Heading 2", WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
style_heading("Heading 3", WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
cap = doc.styles["Caption"]
cap.font.name = FONT; cap.font.size = Pt(10); cap.font.italic = True
cap.font.bold = False; cap.font.color.rgb = BLACK
cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(4); cap.paragraph_format.space_after = Pt(10)

# ------------------------------------------------------------- helpers
def set_margins(section):
    section.left_margin = Cm(4); section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)

set_margins(doc.sections[0])

def add_field(paragraph, instr, placeholder=""):
    r = paragraph.add_run(); fb = OxmlElement("w:fldChar")
    fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    r2 = paragraph.add_run(); it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve"); it.text = instr; r2._r.append(it)
    r3 = paragraph.add_run(); fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate"); r3._r.append(fs)
    if placeholder:
        paragraph.add_run(placeholder)
    r5 = paragraph.add_run(); fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end"); r5._r.append(fe)

def footer_page_number(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in list(p.runs):
        run.text = ""
    add_field(p, " PAGE ", "1")
    for r in p.runs:
        r.font.name = FONT; r.font.size = Pt(11)

def set_pgnum(section, fmt=None, start=None):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); sectPr.append(pg)
    if fmt: pg.set(qn("w:fmt"), fmt)
    if start is not None: pg.set(qn("w:start"), str(start))

_CAPS = [
    ("atención primaria de salud", "Atención Primaria de Salud"),
    ("Atención primaria de salud", "Atención Primaria de Salud"),
    ("Atención Primaria de salud", "Atención Primaria de Salud"),
    ("atención primaria", "Atención Primaria"),
    ("Atención primaria", "Atención Primaria"),
    ("atención farmacéutica", "Atención Farmacéutica"),
    ("Atención farmacéutica", "Atención Farmacéutica"),
    ("químicos farmacéuticos", "Químicos Farmacéuticos"),
    ("Químicos farmacéuticos", "Químicos Farmacéuticos"),
    ("químico farmacéutico", "Químico Farmacéutico"),
    ("Químico farmacéutico", "Químico Farmacéutico"),
]
def cap_terms(t):
    for a, b in _CAPS:
        t = t.replace(a, b)
    return t

def para(text, bold=False, italic=False, size=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=None):
    text = cap_terms(text)
    p = doc.add_paragraph()
    run = p.add_run(text); run.bold = bold; run.italic = italic
    run.font.name = FONT; run.font.size = Pt(size or 12)
    p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, style="List Bullet"):
    text = cap_terms(text)
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for r in p.runs:
        r.font.name = FONT; r.font.size = Pt(12)
    return p

def numbered(text):
    return bullet(text, style="List Number")

def heading(text, level=1):
    return doc.add_heading(cap_terms(text), level=level)

def caption(label, text):
    p = doc.add_paragraph(style="Caption")
    p.add_run(f"{label} ")
    add_field(p, f" SEQ {label} \\* ARABIC ", "1")
    p.add_run(f". {text}")
    return p

def add_image(name, width_cm):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(f"{FIG}/{name}", width=Cm(width_cm))
    return p

def defterm(term, definition):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(cap_terms(f"{term}: ")); r.bold = True; r.font.name = FONT; r.font.size = Pt(12)
    r2 = p.add_run(cap_terms(definition)); r2.font.name = FONT; r2.font.size = Pt(12)
    return p

def make_table(headers, rows, col_widths=None, font_size=10, header_fill="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.name = FONT; run.font.size = Pt(font_size)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), header_fill)
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].text = ""
            run = cells[i].paragraphs[0].add_run(cap_terms(val))
            run.font.name = FONT; run.font.size = Pt(font_size)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
            cells[i].paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

def centered(text, size=12, bold=False, italic=False, space_after=6, space_before=0,
             color=BLACK):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = color
    return p

# =====================================================================
# PORTADA
# =====================================================================
_logo_p = doc.add_paragraph(); _logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
_logo_p.paragraph_format.space_before = Pt(6); _logo_p.paragraph_format.space_after = Pt(18)
_logo_p.add_run().add_picture(f"{FIG}/logo_uss.jpeg", width=Cm(7.5))
centered("FACULTAD DE CIENCIAS", size=14, bold=True)
centered("ESCUELA DE QUÍMICA Y FARMACIA", size=14, bold=True)
centered("SEDE CONCEPCIÓN", size=14, bold=True, space_after=54)
centered("Informe de Internado en Farmacia Asistencial y Atención Primaria de Salud",
         size=13, bold=True, space_after=48)
centered("Lugar de Internado", size=12, space_after=2)
centered("Centro de Salud Familiar Villa Nonguén, Concepción", size=13, bold=True,
         space_after=44)
centered("Estudiante: Bastián Alonso Espinoza Palma", size=12, space_after=6)
centered("Docente Tutor: Q.F. de la Unidad de Farmacia, CESFAM Villa Nonguén",
         size=12, space_after=6)
centered("Docente Supervisor: Escuela de Química y Farmacia, USS", size=12,
         space_after=48)
centered("Concepción, Chile", size=12, space_after=2)
centered("2026", size=12)

# =====================================================================
# SECCION 2: preliminares (romanos)
# =====================================================================
doc.add_section(WD_SECTION.NEW_PAGE)
sec1 = doc.sections[1]; set_margins(sec1)
doc.sections[0].footer.is_linked_to_previous = False
doc.sections[0].footer.paragraphs[0].text = ""
footer_page_number(sec1); set_pgnum(sec1, fmt="lowerRoman", start=2)

heading("TABLA DE CONTENIDOS")
p = doc.add_paragraph()
add_field(p, ' TOC \\o "1-3" \\h \\z \\u ',
          "Actualice este campo en Word: seleccione todo (Ctrl+E) y presione F9.")

doc.add_page_break()
heading("ÍNDICE DE TABLAS")
p = doc.add_paragraph()
add_field(p, ' TOC \\h \\z \\c "Tabla" ', "Actualice este campo en Word (F9).")
heading("ÍNDICE DE FIGURAS")
p = doc.add_paragraph()
add_field(p, ' TOC \\h \\z \\c "Figura" ', "Actualice este campo en Word (F9).")

doc.add_page_break()
heading("ABREVIATURAS")
abrev = [
    ("APS", "Atención Primaria de Salud"),
    ("CESFAM", "Centro de Salud Familiar"),
    ("CENABAST", "Central de Abastecimiento del Sistema Nacional de Servicios de Salud"),
    ("CMO", "Capacidad, Motivación y Oportunidad (modelo de Atención Farmacéutica)"),
    ("COSADES", "Corporación de Salud y Desarrollo Social"),
    ("DM2", "Diabetes Mellitus tipo 2"),
    ("DLP", "Dislipidemia"),
    ("FOFAR", "Fondo de Farmacia"),
    ("GES", "Garantías Explícitas en Salud"),
    ("HTA", "Hipertensión Arterial"),
    ("MINSAL", "Ministerio de Salud de Chile"),
    ("PRM", "Problemas Relacionados con los Medicamentos"),
    ("PSCV", "Programa de Salud Cardiovascular"),
    ("Q.F.", "Químico Farmacéutico"),
    ("RAM", "Reacción Adversa a Medicamentos"),
    ("RNM", "Resultados Negativos asociados a la Medicación"),
    ("TENS", "Técnico en Enfermería de Nivel Superior"),
    ("URM", "Uso Racional de Medicamentos"),
]
for a, d in abrev:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{a}: "); r.bold = True; r.font.name = FONT; r.font.size = Pt(12)
    r2 = p.add_run(d); r2.font.name = FONT; r2.font.size = Pt(12)

doc.add_page_break()
heading("RESUMEN")
para("El presente informe da cuenta del Internado en Farmacia Asistencial y Atención "
     "Primaria de Salud realizado en la unidad de farmacia del CESFAM Villa Nonguén, "
     "entre el 11 de mayo y el 10 de julio de 2026. Durante nueve semanas el interno se "
     "integró al equipo de farmacia y participó en los procesos de recepción, "
     "almacenamiento, gestión de stock, fraccionamiento, reenvasado y dispensación de "
     "medicamentos, además de actividades de atención farmacéutica, visitas "
     "domiciliarias y educación sanitaria orientadas al uso racional de los "
     "medicamentos. En paralelo se desarrolló el seminario de título, consistente en el "
     "diseño de un protocolo de atención farmacéutica domiciliaria basado en el Modelo "
     "CMO para los pacientes del Programa de Salud Cardiovascular del centro. El informe "
     "describe el establecimiento y su unidad de farmacia, detalla las actividades "
     "realizadas mediante una carta Gantt y su descripción, presenta el desarrollo del "
     "seminario de título y discute el cumplimiento de los objetivos a la luz de la "
     "evidencia, junto con las fortalezas, debilidades y propuestas de mejora "
     "identificadas durante la experiencia.")
_pk = doc.add_paragraph(); _pk.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
_pk.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
_rk = _pk.add_run("Palabras clave: "); _rk.bold = True; _rk.font.name = FONT; _rk.font.size = Pt(12)
_rk2 = _pk.add_run("atención farmacéutica; Atención Primaria de Salud; Modelo CMO; "
                   "adherencia; uso racional de medicamentos.")
_rk2.font.name = FONT; _rk2.font.size = Pt(12)
_rk2.text = cap_terms(_rk2.text)

# =====================================================================
# SECCION 3: cuerpo (arabigos)
# =====================================================================
doc.add_section(WD_SECTION.NEW_PAGE)
sec2 = doc.sections[2]; set_margins(sec2)
footer_page_number(sec2); set_pgnum(sec2, fmt="decimal", start=1)

# ---------------------------------------------------------- 1. INTRODUCCION
heading("1. INTRODUCCIÓN")
para("El presente informe reúne la experiencia desarrollada durante el Internado en "
     "Farmacia Asistencial y Atención Primaria de Salud, realizado en la unidad de "
     "farmacia del Centro de Salud Familiar (CESFAM) Villa Nonguén, en la comuna de "
     "Concepción, Región del Biobío, entre el 11 de mayo y el 10 de julio de 2026. El "
     "internado corresponde a una actividad obligatoria del plan de estudios de la "
     "carrera de Química y Farmacia, cuyo propósito es que el estudiante adquiera "
     "experiencia directa del quehacer profesional del Químico Farmacéutico en un "
     "establecimiento de ejercicio real, bajo la supervisión de un profesional del "
     "área. La Atención Primaria constituye la puerta de entrada al sistema público de "
     "salud y, dentro de ella, la farmacia cumple un rol central en el acceso a los "
     "medicamentos y en el uso seguro y racional de la farmacoterapia por parte de la "
     "población (Ministerio de Salud de Chile, 2018b).")

heading("1.1 Aspectos generales del centro de internado", level=2)
para("El CESFAM Villa Nonguén es un establecimiento del nivel primario de atención, "
     "ubicado en Río Loa 1397, en el sector de Nonguén, en la periferia oriente de "
     "Concepción, e integrado a la red asistencial del Servicio de Salud Concepción. "
     "Su origen se remonta a 1987, cuando profesionales de la entonces Octava Región, "
     "con apoyo de la cooperación internacional, elaboraron un proyecto de salud "
     "comunitaria para el sector. El establecimiento inició sus funciones asistenciales "
     "en 1991 y, en 1993, fue seleccionado por el Ministerio de Salud como uno de los "
     "centros pioneros en implementar el Modelo de Salud Familiar y el financiamiento "
     "per cápita en el país (Diario Concepción, 2017). Desde entonces funciona bajo la "
     "modalidad de administración delegada, a cargo de la Corporación de Salud y "
     "Desarrollo Social (COSADES), lo que le permite formar parte de la red pública "
     "manteniendo cierta autonomía en su gestión (COSADES, s.f.).")
para("El centro atiende a una población inscrita validada cercana a las 16.760 "
     "personas, organizada en torno a unas 5.560 familias, con un promedio de 3,1 "
     "integrantes por grupo familiar, y entrega prestaciones a todos los grupos del "
     "ciclo vital, desde la gestación hasta la persona mayor. Para ordenar la atención, "
     "el territorio se divide en cuatro sectores identificados por colores (verde, "
     "azul, café y blanco), lo que facilita la planificación de los controles, las "
     "visitas domiciliarias y las actividades comunitarias según el lugar de residencia "
     "de los usuarios. La dotación total del establecimiento alcanza los 122 "
     "funcionarios distribuidos en distintas profesiones y oficios, y contempla una "
     "jornada diurna y otra vespertina. La Tabla 1 resume los principales antecedentes "
     "del centro y la Tabla 2 detalla su dotación de personal.")
caption("Tabla", "Antecedentes generales del CESFAM Villa Nonguén.")
make_table(["Antecedente", "Detalle"], [
    ["Dependencia administrativa", "Administración delegada, COSADES, Servicio de Salud Concepción"],
    ["Nivel de atención", "Primario (Atención Primaria de Salud)"],
    ["Dirección", "Río Loa 1397, Villa Nonguén, Concepción, Región del Biobío"],
    ["Población inscrita validada", "16.760 personas aproximadamente"],
    ["Familias registradas", "5.560 (promedio de 3,1 integrantes por familia)"],
    ["Sectorización", "Cuatro sectores: verde, azul, café y blanco"],
    ["Dotación total", "122 funcionarios (jornada diurna y vespertina)"],
], col_widths=[5.5, 9.5])
caption("Tabla", "Dotación de personal del CESFAM Villa Nonguén según estamento.")
make_table(["Estamento", "N.º", "Estamento", "N.º"], [
    ["Médico", "14", "Enfermera/o", "11"],
    ["Odontólogo", "7", "Matrona", "8"],
    ["Químico Farmacéutico", "3", "Kinesiólogo", "3"],
    ["Nutricionista", "3", "Psicólogo", "6"],
    ["Trabajador social", "5", "Podólogo", "1"],
    ["Técnico en enfermería", "25", "Secretaria", "11"],
    ["Auxiliar de servicios", "8", "Otros (apoyo administrativo)", "10"],
], col_widths=[5.6, 1.9, 5.6, 1.9])
para("Nota: La dotación total corresponde a 122 funcionarios. Elaboración a partir de "
     "los antecedentes entregados por el establecimiento.", size=10, italic=True,
     space_after=10)

heading("1.2 Organigrama y distribución de la unidad de farmacia", level=2)
para("La unidad de farmacia se organiza bajo la responsabilidad de un químico "
     "farmacéutico que ejerce la dirección técnica y responde por el funcionamiento "
     "general de la unidad, el cumplimiento normativo y la gestión del arsenal "
     "farmacoterapéutico. Un segundo Químico Farmacéutico se dedica de manera "
     "preferente a la Atención Farmacéutica, tanto en box como en las visitas "
     "domiciliarias del Programa de Salud Cardiovascular, y un equipo permanente de "
     "cuatro técnicos en enfermería de nivel superior (TENS) apoya las tareas de "
     "recepción, almacenamiento, fraccionamiento, reenvasado y dispensación. El interno "
     "de Química y Farmacia se integró a este equipo durante toda la rotación. La "
     "Figura 1 presenta el organigrama de la unidad, con los Químicos Farmacéuticos y "
     "los técnicos que componen el equipo.")
add_image("fig_organigrama.png", 15.0)
caption("Figura", "Organigrama de la unidad de farmacia del CESFAM Villa Nonguén.")

heading("1.3 Servicios que brinda el centro", level=2)
para("El CESFAM Villa Nonguén entrega una atención integral y continua a lo largo del "
     "curso de vida, con acciones de promoción, prevención, tratamiento y "
     "rehabilitación. Su cartera de prestaciones incluye, entre otros, el Programa de "
     "Salud de la Mujer, el Programa Nacional de Salud de la Infancia, el Programa de "
     "Salud Integral del Adolescente, el Programa de Salud Cardiovascular, el Programa "
     "de Salud Mental, el Programa Nacional de Salud Integral de Personas Mayores, el "
     "Programa de Salud Familiar y el Programa VIH/ITS. A ello se suman prestaciones de "
     "procedimientos, vacunatorio, toma de muestras, salud dental, controles de salud, "
     "exámenes preventivos y atención domiciliaria, junto con la entrega de alimentos "
     "de los programas PNAC y PACAM.")
para("La unidad de farmacia participa de manera transversal en estos programas a "
     "través de la entrega de medicamentos, el fraccionamiento y reenvasado según las "
     "indicaciones, el control del arsenal y de los medicamentos sujetos a control "
     "legal, y la Atención Farmacéutica orientada al uso seguro y racional de la "
     "terapia. Dentro del Programa de Salud Cardiovascular, el Químico Farmacéutico "
     "acompaña a los usuarios en box y en su domicilio, con seguimiento "
     "farmacoterapéutico, educación y revisión de la medicación, y participa en la "
     "gestión del Fondo de Farmacia (FOFAR), que financia su presencia en la atención "
     "primaria y el acceso a los medicamentos para la hipertensión, la diabetes y la "
     "dislipidemia (Ministerio de Salud de Chile, 2020).")

heading("1.4 Marco teórico", level=2)
para("La Atención Farmacéutica se entiende hoy como la participación del farmacéutico "
     "en el cuidado de la persona para aprovechar mejor sus medicamentos y mejorar sus "
     "resultados en salud. Bonal y colaboradores la describen como el proceso por el "
     "cual el farmacéutico coopera con el paciente y con el resto del equipo en el "
     "diseño, la ejecución y el seguimiento de un plan terapéutico orientado a metas "
     "concretas (Bonal et al., 2003). Este enfoque supone un cambio respecto del modelo "
     "clásico, centrado en entregar el medicamento, hacia otro que se preocupa de lo "
     "que ocurre con ese medicamento una vez que el paciente lo lleva a su casa.")
para("En la Atención Primaria chilena, la incorporación del Químico Farmacéutico a "
     "los equipos ha permitido acercar esta labor a la comunidad, con especial "
     "relevancia en las enfermedades crónicas. La Encuesta Nacional de Salud estimó que "
     "cerca del 27,6% de la población tiene sospecha de hipertensión, un 12,3% diabetes "
     "tipo 2 y alrededor del 60% algún grado de dislipidemia (Ministerio de Salud de "
     "Chile, 2017a), condiciones que se atienden principalmente a través del Programa "
     "de Salud Cardiovascular. Ahora bien, que un tratamiento esté indicado no asegura "
     "que el paciente lo tome como corresponde. La adherencia, entendida como el grado "
     "en que la persona sigue las indicaciones acordadas con el equipo de salud, es uno "
     "de los factores que más influye en el control de estas enfermedades y suele ser "
     "baja en quienes usan varios medicamentos a la vez (Haynes et al., 1979; Martín "
     "Alfonso, 2004). Para medirla de forma sencilla se utilizan instrumentos como el "
     "test de Morisky-Green-Levine (Morisky et al., 1986), y para clasificar los "
     "problemas detectados se recurre al Tercer Consenso de Granada, que ordena los "
     "resultados negativos asociados a la medicación según si el tratamiento es "
     "necesario, efectivo y seguro (Comité de Consenso de Granada, 2007).")
para("El marco que hoy se usa como referencia para modernizar la atención "
     "farmacéutica es el Modelo CMO, sigla de Capacidad, Motivación y Oportunidad, "
     "desarrollado en España y validado por su sociedad de farmacia hospitalaria "
     "(Calleja Hernández & Morillo Verdugo, 2016). Su valor está en que ordena la "
     "atención en torno a tres ideas que pueden llevarse al domicilio del paciente. La "
     "Capacidad tiene que ver con priorizar, es decir, concentrar el tiempo en los "
     "pacientes que más lo necesitan según criterios de riesgo y complejidad (Manzano "
     "García & Morillo Verdugo, 2018). La Motivación pone el foco en la relación con el "
     "paciente y en su adherencia, trabajando mediante la entrevista motivacional las "
     "creencias y dificultades que llevan a no tomar bien el tratamiento. La "
     "Oportunidad, por último, busca llegar al paciente más allá del box, en su casa o "
     "por teléfono, para resolver a tiempo lo que vaya surgiendo. Estos tres pilares "
     "son los que dan forma al protocolo desarrollado como seminario de título durante "
     "el internado, descrito en la sección 2.4.")

heading("1.5 Terminología y definiciones", level=2)
para("Para facilitar la lectura del informe se precisan a continuación algunos "
     "términos propios del trabajo desarrollado en la unidad de farmacia.")
defterm("Dispensación informada", "acto profesional de entrega del medicamento "
        "acompañado de la información necesaria para su uso correcto, seguro y efectivo.")
defterm("Fraccionamiento y reenvasado", "proceso mediante el cual se acondiciona el "
        "medicamento en la cantidad que corresponde a cada paciente, rotulando el envase "
        "con el nombre, la indicación, la cantidad, la fecha de vencimiento y la serie o "
        "lote, de manera que quede asegurada la trazabilidad.")
defterm("Seguimiento farmacoterapéutico", "práctica clínica en la que el farmacéutico "
        "detecta, previene y resuelve los problemas relacionados con los medicamentos de "
        "forma continua y documentada (Ministerio de Salud de Chile, 2018a).")
defterm("Conciliación de la medicación", "comparación entre la medicación que el "
        "paciente tiene indicada y la que realmente utiliza, para acordar los cambios "
        "que correspondan (Delgado et al., 2007).")
defterm("Problemas relacionados con los medicamentos (PRM)", "situaciones que, en el "
        "proceso de uso de los medicamentos, pueden causar o causan un resultado "
        "negativo asociado a la medicación.")
defterm("Resultados negativos asociados a la medicación (RNM)", "resultados en la "
        "salud del paciente que no se corresponden con los objetivos de la "
        "farmacoterapia (Comité de Consenso de Granada, 2007).")

heading("1.6 Objetivos", level=2)
heading("1.6.1 Objetivo general", level=3)
para("Desarrollar las competencias profesionales del Químico Farmacéutico en el "
     "ámbito de la farmacia asistencial y la Atención Primaria de Salud durante el "
     "internado en el CESFAM Villa Nonguén.")
heading("1.6.2 Objetivos específicos", level=3)
numbered("Reconocer la estructura organizacional y las funciones del equipo de la "
         "unidad de farmacia del CESFAM Villa Nonguén.")
numbered("Participar en los procesos técnicos de recepción, almacenamiento, "
         "fraccionamiento, reenvasado y dispensación de medicamentos de la unidad de "
         "farmacia.")
numbered("Ejecutar actividades de Atención Farmacéutica, visitas domiciliarias y "
         "educación sanitaria orientadas al uso racional de los medicamentos.")
numbered("Diseñar un protocolo de Atención Farmacéutica domiciliaria basado en el "
         "Modelo CMO para el Programa de Salud Cardiovascular del CESFAM Villa Nonguén.")
para("Cada uno de estos objetivos se desarrolla de forma explícita en el informe: los "
     "tres primeros a través de las actividades descritas en la sección 2.3, y el "
     "cuarto a través del seminario de título presentado en la sección 2.4. Su grado "
     "de cumplimiento se evalúa en la sección 3.1.", space_after=8)

# ---------------------------------------------------- 2. ACTIVIDADES REALIZADAS
heading("2. ACTIVIDADES REALIZADAS")
para("El internado se organizó en nueve semanas, entre el 11 de mayo y el 10 de julio "
     "de 2026. Durante ese periodo el interno rotó por las distintas tareas de la "
     "unidad de farmacia y, en paralelo, desarrolló el seminario de título. Las "
     "actividades realizadas fueron las siguientes:")
acts_list = [
    "Inducción y conocimiento del centro y de la unidad de farmacia.",
    "Recepción, almacenamiento y gestión del stock de medicamentos.",
    "Fraccionamiento y reenvasado de medicamentos.",
    "Dispensación informada de medicamentos.",
    "Control de indicadores y registros de despacho.",
    "Atención Farmacéutica y visitas domiciliarias.",
    "Educación sanitaria y promoción del uso racional de medicamentos.",
    "Inventario y ordenamiento de bodega.",
    "Apoyo en las supervisiones del Servicio de Salud.",
    "Vinculación con el medio.",
    "Seminario de título: diseño de un protocolo de Atención Farmacéutica domiciliaria.",
]
for a in acts_list:
    numbered(a)

heading("2.1 Cronograma de actividades planificadas", level=2)
para("Al inicio del internado se acordó con el equipo de farmacia un plan de trabajo "
     "que ordenó la rotación por las distintas tareas de la unidad y reservó tiempo "
     "para el desarrollo del seminario de título. La Figura 2 presenta la carta Gantt "
     "con las actividades planificadas a lo largo de las nueve semanas.")
add_image("fig_gantt_planificada.png", 15.5)
caption("Figura", "Carta Gantt de actividades planificadas del internado.")

heading("2.2 Cronograma de actividades desarrolladas", level=2)
para("La Figura 3 muestra las actividades efectivamente desarrolladas, con detalle por "
     "día y por semana. En términos generales, la ejecución siguió lo planificado, con "
     "algunos ajustes propios de la dinámica del establecimiento. El fraccionamiento, "
     "la dispensación y la Atención Farmacéutica se extendieron durante casi todo el "
     "periodo, ya que constituyen tareas habituales de la unidad, y se sumaron "
     "actividades no previstas al inicio, como el apoyo durante las supervisiones del "
     "Servicio de Salud y una situación de vinculación con el medio surgida en terreno. "
     "La Tabla 3 complementa la figura con el detalle diario de lo realizado.")
add_image("fig_gantt_desarrollada.png", 15.0)
caption("Figura", "Carta Gantt de actividades desarrolladas, con detalle por día y semana.")

caption("Tabla", "Cronograma detallado de actividades por día y semana.")
make_table(
    ["Semana", "Lunes", "Martes", "Miércoles", "Jueves", "Viernes"],
    [
        ["S1\n11-15 may", "Inducción y recorrido de la unidad",
         "Fraccionamiento y reenvasado; cambio de vildagliptina",
         "Control de vencimientos; indicador de despacho por sector",
         "Envasado; protocolos de emergencia; llegada de pedido; Atención Farmacéutica",
         "Seminario: revisión bibliográfica y Modelo CMO"],
        ["S2\n18-22 may", "Protocolo de AF y entrega a domicilio",
         "Seminario: objetivos; bolsas y pedido", "Fraccionamiento y reenvasado",
         "Dispensación", "Gestión de stock y dispensación"],
        ["S3\n25-29 may", "Entrega a domicilio; apoyo a adulto mayor",
         "Entrega a domicilio; objetivos del seminario",
         "Clase de bibliografía (Mendeley)",
         "Dispensación; ubicación de medicamentos", "Dispensación"],
        ["S4\n01-05 jun", "Préstamos por vencimiento", "Plantilla de despacho de mayo",
         "Gestión de stock y reportes",
         "Supervisión del Servicio de Salud; inducción; metodología del seminario",
         "Indicador y Atención Farmacéutica"],
        ["S5\n08-12 jun", "Rutina de farmacia y dispensación",
         "Inventario y ordenamiento de bodega",
         "Inspección del Servicio de Salud; avance del seminario",
         "Fraccionamiento y dispensación", "Bolsas fraccionadas; visita domiciliaria"],
        ["S6\n15-19 jun", "Llega docente supervisora; envío del avance",
         "Despacho a domicilio y farmacia", "Despacho a domicilio y farmacia",
         "Despacho a domicilio y farmacia", "Atención Farmacéutica"],
        ["S7\n22-26 jun", "Tareas generales de farmacia", "Tareas generales de farmacia",
         "Fraccionamiento y farmacia", "Tareas generales de farmacia",
         "Taller de uso racional de medicamentos"],
        ["S8\n29 jun-03 jul", "Rutina de farmacia", "Jornada vespertina; cierre de farmacia",
         "Rutina de farmacia", "Salida a terreno para dejar medicamentos",
         "Avance del seminario"],
        ["S9\n06-10 jul", "Entrega a domicilio", "Entrega a domicilio",
         "Entrega a domicilio", "Farmacia y fraccionamiento",
         "Atención Farmacéutica con Q.F. especialista"],
    ],
    col_widths=[2.0, 2.6, 2.6, 2.6, 2.6, 2.6], font_size=8)
para("Respecto de las actividades no realizadas, la participación en el Comité de "
     "Farmacia y Terapéutica no pudo concretarse porque el comité no sesionó durante el "
     "periodo de la rotación, de modo que su contenido se abordó de manera indirecta a "
     "través de la revisión de los criterios de selección del arsenal y de los cambios "
     "de medicamentos informados por el Servicio de Salud. Las actividades formales de "
     "farmacovigilancia se limitaron a conocer el circuito de notificación, sin que se "
     "presentara un caso que ameritara notificar durante la estadía.")

heading("2.3 Descripción de las actividades realizadas", level=2)

heading("2.3.1 Inducción y conocimiento del centro y de la unidad de farmacia", level=3)
para("Durante la primera semana se realizó la presentación al equipo de trabajo y se "
     "recorrieron las dependencias del CESFAM y de la unidad de farmacia. El equipo "
     "explicó la organización de la unidad, la ubicación física de los medicamentos, el "
     "flujo de trabajo diario y el uso de los sistemas de registro. Se revisaron los "
     "protocolos internos del establecimiento, incluidos los de emergencia, como el "
     "código rojo y la reanimación cardiopulmonar ante un paro cardiorrespiratorio, y "
     "el manejo inicial de situaciones críticas como la reacción alérgica, la "
     "hemorragia masiva y la convulsión. También se conocieron los programas "
     "computacionales de apoyo y la ficha clínica electrónica del establecimiento "
     "(SINET Sur), utilizada para consultar antecedentes y dejar constancia de las "
     "atenciones. Esta actividad responde al primer objetivo específico.")

heading("2.3.2 Recepción, almacenamiento y gestión del stock", level=3)
para("Se participó en la recepción y el almacenamiento de los medicamentos que llegan "
     "a la unidad, que se adquieren principalmente a través de la Central de "
     "Abastecimiento (CENABAST) y que se ordenan resguardando las condiciones de "
     "conservación. Se colaboró en la gestión del stock, con especial atención al "
     "control de los medicamentos próximos a vencer, que se marcaban y se controlaban "
     "para evitar pérdidas. Cuando un medicamento estaba por vencer y no alcanzaba a "
     "utilizarse, se gestionaban préstamos o canjes con otros centros de la red, y las "
     "existencias que no se podían aprovechar se reportaban y se redistribuían hacia "
     "otros CESFAM, de manera de dar un uso eficiente a los recursos y reducir las "
     "mermas. Se conoció además el proceso de programación anual de las necesidades de "
     "medicamentos, que se planifica según la demanda y las metas del establecimiento. "
     "Esta actividad responde al segundo objetivo específico.")

heading("2.3.3 Fraccionamiento y reenvasado de medicamentos", level=3)
para("El fraccionamiento y el reenvasado fueron tareas frecuentes a lo largo del "
     "internado. El trabajo consistió en acondicionar los medicamentos en la cantidad "
     "correspondiente a cada paciente y preparar las bolsas de tratamiento, rotulando "
     "cada envase con el nombre del medicamento, la indicación, la cantidad, la fecha "
     "de vencimiento y la serie o lote, de manera de asegurar la trazabilidad del "
     "producto hasta el usuario. Este proceso permitió comprender la importancia del "
     "rotulado correcto y del resguardo de la información para la seguridad del "
     "paciente, sobre todo en personas con polifarmacia que retiran varios medicamentos "
     "a la vez. Responde al segundo objetivo específico.")

heading("2.3.4 Dispensación informada de medicamentos", level=3)
para("Se apoyó la dispensación de medicamentos a los usuarios, con lo que se "
     "reforzaron el reconocimiento de los principios activos, las presentaciones y su "
     "ubicación dentro de la unidad, y las buenas prácticas de entrega. La dispensación "
     "se acompañó de la información necesaria para el uso correcto de la terapia, "
     "verificando la indicación y resolviendo dudas de los pacientes. También se "
     "conocieron los criterios de manejo de los medicamentos sujetos a control legal y "
     "la forma de registrar su entrega. Responde al segundo objetivo específico.")

heading("2.3.5 Atención Farmacéutica y visitas domiciliarias", level=3)
para("La Atención Farmacéutica fue una de las actividades más significativas del "
     "internado. Se acompañó al Químico Farmacéutico en la atención de los usuarios del "
     "Programa de Salud Cardiovascular, tanto en box como en el domicilio, y se "
     "participó en la entrega de medicamentos a domicilio para pacientes con "
     "dificultades para acudir al centro. En estas visitas se revisaba la medicación "
     "que el paciente realmente utilizaba, se conciliaba con la que tenía indicada, se "
     "resolvían dudas sobre horarios y formas de administración, y se reforzaba la "
     "adherencia. Durante una jornada vespertina se observó, además, el funcionamiento "
     "de la farmacia en el horario de la tarde y el procedimiento de cierre de la "
     "unidad. Esta actividad responde al tercer objetivo específico.")

heading("2.3.6 Educación sanitaria y promoción del uso racional de medicamentos", level=3)
para("Se participó en un taller grupal de uso racional de medicamentos dirigido a los "
     "usuarios del CESFAM. La actividad se realizó junto al Químico Farmacéutico y "
     "consistió en conversar con los pacientes, entregar material educativo en forma de "
     "trípticos y desarrollar una dinámica participativa con paletas de respuesta, en "
     "la que los asistentes respondían situaciones sobre el uso correcto de los "
     "medicamentos y recibían pequeños incentivos por participar. La instancia "
     "favoreció la educación sanitaria y acercó el rol del farmacéutico a la comunidad. "
     "Responde al tercer objetivo específico.")

heading("2.3.7 Control de indicadores, inventario y supervisiones", level=3)
para("Se colaboró en el registro y el control de los indicadores de la unidad. Entre "
     "ellos, se trabajó en el indicador de porcentaje de despacho de medicamentos "
     "realizado según el protocolo de Atención Farmacéutica, desagregado por sector, y "
     "se completó la plantilla mensual de despacho correspondiente al mes de mayo. "
     "También se participó en las tareas de inventario y en el ordenamiento de la "
     "bodega, cuidando que las existencias quedaran correctamente ubicadas y que no "
     "permanecieran cajas en el piso. Durante el internado el establecimiento recibió "
     "supervisiones e inspecciones del Servicio de Salud, en cuya preparación se "
     "acompañó al equipo, lo que permitió conocer los aspectos que se evalúan en la "
     "unidad de farmacia y la relevancia de mantener los registros al día. Además, en "
     "una salida a terreno se prestó apoyo a un adulto mayor para regresar a su hogar, "
     "situación que reflejó el componente humano y comunitario del trabajo en la "
     "Atención Primaria.")

heading("2.4 Seminario de título", level=2)
para("El seminario de título se desarrolló en paralelo a las actividades de farmacia y "
     "corresponde al cuarto objetivo específico del internado. Consistió en el diseño "
     "de un protocolo de Atención Farmacéutica domiciliaria basado en el Modelo CMO "
     "para los pacientes del Programa de Salud Cardiovascular del CESFAM Villa Nonguén, "
     "bajo la guía del profesor Diego Ignacio Jorquera Pereira y con el apoyo del "
     "Químico Farmacéutico tutor del centro. A continuación se describe cómo se realizó.")

heading("2.4.1 Objetivos del seminario", level=3)
para("El objetivo general del seminario fue diseñar un protocolo de atención "
     "farmacéutica domiciliaria basado en el Modelo CMO para los pacientes del Programa "
     "de Salud Cardiovascular del CESFAM Villa Nonguén. Para alcanzarlo se plantearon "
     "tres objetivos específicos: identificar las brechas del programa de visitas "
     "domiciliarias que se realizaba en el centro; definir el flujograma de la visita y "
     "los criterios para priorizar a los pacientes según los pilares del Modelo CMO; y "
     "diseñar los instrumentos y herramientas clínicas necesarias para aplicar el "
     "protocolo, referidas a la adherencia, a los PRM y RNM y al registro de la "
     "atención.")

heading("2.4.2 Tipo de estudio y desarrollo", level=3)
para("El seminario correspondió a un trabajo de desarrollo metodológico. No se realizó "
     "un experimento ni se recogieron datos clínicos de pacientes con fines de "
     "investigación, sino que se diseñó un protocolo a partir de la revisión de la "
     "literatura, del análisis de la normativa del Ministerio de Salud y de la "
     "adaptación de herramientas ya utilizadas en la Atención Primaria, ajustadas a lo "
     "observado en el propio centro. El trabajo se desarrolló durante los meses del "
     "internado, en el CESFAM Villa Nonguén y su territorio. La revisión bibliográfica "
     "se realizó en bases de datos como PubMed, SciELO y Google Scholar, junto con la "
     "normativa y las guías del Ministerio de Salud, y las referencias se ordenaron con "
     "el gestor bibliográfico Mendeley. Como apoyo para la información de medicamentos "
     "se consultaron fuentes de referencia como UpToDate, Medscape y Drugs.com.")

heading("2.4.3 Diagnóstico de brechas, población y criterios", level=3)
para("En una primera etapa se diagnosticaron las brechas del programa de visitas "
     "domiciliarias vigente, comparando la práctica del centro con las orientaciones "
     "técnicas del Programa de Salud Cardiovascular y del Fondo de Farmacia, la Guía de "
     "Atención Farmacéutica y Seguimiento Farmacoterapéutico en APS (Ministerio de "
     "Salud de Chile, 2018a) y los tres pilares del Modelo CMO. Como referencia se "
     "utilizó también la evaluación previa del mismo programa, que había seguido a "
     "veinte usuarios y encontrado que todos tenían al menos un problema relacionado "
     "con sus medicamentos y que solo el 15% era adherente al inicio del seguimiento "
     "(Faúndez Navarrete, 2020). La Tabla 4 resume las brechas detectadas. El protocolo "
     "se orientó a las personas adultas y mayores del programa, con prioridad en "
     "quienes toman muchos medicamentos, dependen de un cuidador, se descompensan con "
     "frecuencia o tienen dificultades para acudir al centro, y se definieron criterios "
     "de inclusión y de exclusión para ordenar el ingreso al seguimiento.")
caption("Tabla", "Brechas del programa de visitas domiciliarias según los pilares del Modelo CMO.")
make_table(
    ["Pilar del Modelo CMO", "Lo que se espera", "Lo observado", "Brecha"],
    [
        ["Capacidad", "Priorizar la visita según el riesgo y la complejidad.",
         "El ingreso depende de la derivación y del tiempo disponible (unas 25 visitas al mes).",
         "No hay un criterio de estratificación de riesgo propio del farmacéutico."],
        ["Motivación", "Abordar la adherencia con entrevista motivacional estructurada.",
         "Muchas derivaciones son por adherencia o desconocimiento del tratamiento.",
         "La educación y el trabajo de barreras dependen de cada visita."],
        ["Oportunidad", "Registro estructurado y continuidad del seguimiento.",
         "Registro en la ficha clínica (SINET Sur) y en hojas de apoyo personales.",
         "Falta un instrumento de registro propio de la Atención Farmacéutica."],
    ],
    col_widths=[3.0, 4.0, 4.0, 4.0], font_size=9)

heading("2.4.4 Instrumentos diseñados y consideraciones éticas", level=3)
para("El protocolo tradujo los tres pilares del Modelo CMO en pasos concretos. La "
     "Capacidad se abordó mediante la priorización de los pacientes y la revisión del "
     "botiquín del hogar para conciliar la medicación (Delgado et al., 2007); la "
     "Motivación, mediante la entrevista motivacional orientada a trabajar las "
     "creencias y barreras de adherencia; y la Oportunidad, mediante la continuidad del "
     "seguimiento por vía telefónica y el refuerzo del vínculo con la farmacia. Para su "
     "aplicación se diseñaron y adaptaron instrumentos sencillos, compatibles con la "
     "rutina del CESFAM: una ficha de conciliación, una ficha de seguimiento "
     "farmacoterapéutico, el test de Morisky-Green-Levine para la adherencia, la "
     "clasificación de los PRM y RNM según el Consenso de Granada, una hoja con el "
     "esquema de horarios y un protocolo de llamada telefónica para agendar la visita "
     "(ver Anexos). Al tratarse del diseño de un protocolo y no de un estudio con "
     "pacientes, el trabajo no contempló intervenciones experimentales ni la "
     "recolección de datos clínicos sensibles con fines de investigación, y se "
     "resguardaron los principios de confidencialidad y de uso responsable de la "
     "información clínica. La propuesta quedó a disposición de los químicos "
     "farmacéuticos del establecimiento para su revisión y eventual incorporación a la "
     "práctica.")

# ---------------------------------------------- 3. DISCUSION Y CONCLUSIONES
heading("3. DISCUSIÓN Y CONCLUSIONES")

heading("3.1 Alcances y cumplimiento de los objetivos", level=2)
para("El internado permitió conocer de manera directa el funcionamiento de una unidad "
     "de farmacia de Atención Primaria y participar en la mayoría de sus procesos, "
     "desde la gestión del arsenal hasta la Atención Farmacéutica en el domicilio. El "
     "primer objetivo específico se cumplió al reconocer la estructura, las funciones "
     "del equipo y los procesos de la unidad, tal como se describe en la sección 1 y en "
     "el organigrama. El segundo se alcanzó mediante la participación sostenida en la "
     "recepción, el almacenamiento, la gestión de stock, el fraccionamiento, el "
     "reenvasado y la dispensación, resguardando la trazabilidad a través del rotulado. "
     "El tercero se desarrolló a través de la Atención Farmacéutica, las visitas "
     "domiciliarias y el taller de uso racional de medicamentos. Y el cuarto se "
     "materializó en el diseño del protocolo de Atención Farmacéutica domiciliaria "
     "basado en el Modelo CMO, presentado como seminario de título. En conjunto, estas "
     "actividades permitieron cumplir el objetivo general del internado.")
para("Los resultados observados en terreno concuerdan con lo descrito en la literatura "
     "reciente. La evidencia muestra que las intervenciones farmacéuticas domiciliarias "
     "mejoran la adherencia y el conocimiento de los pacientes sobre su tratamiento, "
     "sobre todo en personas mayores y con polifarmacia (Ahn et al., 2024), y que el "
     "acompañamiento sostenido en el tiempo produce mejores resultados de adherencia en "
     "pacientes crónicos (Lambert et al., 2024). En el ámbito cardiovascular, la falta "
     "de adherencia sigue siendo una de las principales barreras para el control de la "
     "hipertensión, la diabetes y la dislipidemia, y las estrategias más eficaces "
     "combinan la comunicación con el paciente, la continuidad de la atención y la "
     "participación del farmacéutico (Espinosa García et al., 2023). Asimismo, las "
     "experiencias de aplicación del Modelo CMO en distintos grupos de pacientes "
     "crónicos han mostrado mejoras en la adherencia, en los objetivos "
     "farmacoterapéuticos y en la experiencia del usuario (Cantillana-Suárez et al., "
     "2021; Sánchez-Yáñez et al., 2023), aunque su adopción todavía enfrenta barreras "
     "organizacionales que conviene anticipar (Álvarez-Díaz et al., 2025; "
     "Morillo-Verdugo et al., 2022). Estos antecedentes respaldan la pertinencia del "
     "protocolo diseñado durante el internado.")

heading("3.2 Fortalezas y debilidades del centro", level=2)
para("Entre las fortalezas del CESFAM Villa Nonguén destacan su condición de centro "
     "pionero en el Modelo de Salud Familiar, un equipo de farmacia consolidado que ya "
     "realiza Atención Farmacéutica en box y en el domicilio, y una cultura de trabajo "
     "orientada a la comunidad. La sectorización del territorio y la existencia de "
     "indicadores de despacho facilitan la organización de la atención y el uso "
     "eficiente de los recursos. Como aspecto por mejorar, se observó que la visita "
     "domiciliaria del farmacéutico se realizaba sin un procedimiento escrito y "
     "estandarizado, y que el registro se apoyaba en hojas personales además de la "
     "ficha clínica electrónica, lo que dificultaba la comparación entre profesionales "
     "y el seguimiento en el tiempo. La carga de trabajo y la disponibilidad de tiempo "
     "también limitaban la frecuencia de las visitas.")

heading("3.3 Sugerencias y propuestas de mejora", level=2)
numbered("Incorporar de manera formal el protocolo de Atención Farmacéutica "
         "domiciliaria basado en el Modelo CMO, con criterios de priorización propios "
         "del farmacéutico, para ordenar el ingreso de los pacientes al seguimiento.")
numbered("Estandarizar el registro de la Atención Farmacéutica mediante fichas de "
         "conciliación y de seguimiento comunes, que luego se traspasen a la ficha "
         "clínica electrónica, para mejorar la trazabilidad y la continuidad.")
numbered("Reforzar el trabajo de la adherencia con herramientas simples y validadas, "
         "como el test de Morisky-Green-Levine, y con la entrevista motivacional.")
numbered("Aprovechar la continuidad telefónica y las instancias educativas grupales "
         "para mantener el vínculo con los pacientes entre una visita y otra.")

heading("3.4 Conclusiones", level=2)
numbered("El internado en el CESFAM Villa Nonguén permitió desarrollar las "
         "competencias del Químico Farmacéutico en el ámbito de la farmacia asistencial "
         "y la Atención Primaria, cumpliendo el objetivo general planteado.")
numbered("Se reconocieron la estructura y los procesos de la unidad de farmacia, y se "
         "participó en las tareas de recepción, almacenamiento, gestión de stock, "
         "fraccionamiento, reenvasado y dispensación, resguardando la trazabilidad y el "
         "uso eficiente de los recursos.")
numbered("Se ejecutaron actividades de Atención Farmacéutica, visitas domiciliarias y "
         "promoción del uso racional de los medicamentos, lo que evidenció el aporte "
         "clínico y comunitario del farmacéutico en este nivel de atención.")
numbered("Se diseñó, como seminario de título, un protocolo de Atención Farmacéutica "
         "domiciliaria basado en el Modelo CMO, que aporta una herramienta concreta "
         "para ordenar, estandarizar y hacer trazable la visita domiciliaria en el "
         "establecimiento.")
numbered("La experiencia confirmó la relevancia de la Atención Primaria como puerta de "
         "entrada al sistema de salud y el valor de un ejercicio profesional basado en "
         "la evidencia, la seguridad del paciente y el trabajo en equipo.")

# --------------------------------------------------------------- BIBLIOGRAFIA
heading("BIBLIOGRAFÍA")
refs = [
    "Ahn, H., Lee, S., Kim, S., Jang, S., & Bae, S. (2024). Effects of pharmacist-led "
    "home visit services and factors influencing medication adherence improvement. PLoS "
    "ONE, 19(11), e0313101.",
    "Álvarez-Díaz, A. M., Morillo-Verdugo, R., & Fernández-Llamazares, C. M. (2025). "
    "Estudio cualitativo sobre la adopción y potenciación del modelo "
    "capacidad-motivación-oportunidad para la Atención Farmacéutica en consultas "
    "externas de farmacia en España. Farmacia Hospitalaria, 49(6), 384-391.",
    "Bonal, J., Alerany, C., Bassons, T., & Gascón, P. (2003). Farmacia clínica y "
    "Atención Farmacéutica. Sociedad Española de Farmacia Hospitalaria.",
    "Calleja Hernández, M. Á., & Morillo Verdugo, R. (Eds.). (2016). El modelo CMO en "
    "consultas externas de farmacia hospitalaria. Sociedad Española de Farmacia "
    "Hospitalaria.",
    "Cantillana-Suárez, M. G., Robustillo-Cortés, M. A., Gutiérrez-Pizarraya, A., & "
    "Morillo-Verdugo, R. (2021). Impact and acceptance of pharmacist-led interventions "
    "during HIV care using the Capacity-Motivation-Opportunity model: The IRAFE study. "
    "European Journal of Hospital Pharmacy, 28(e1), e157-e163.",
    "Comité de Consenso de Granada. (2007). Tercer Consenso de Granada sobre problemas "
    "relacionados con medicamentos (PRM) y resultados negativos asociados a la "
    "medicación (RNM). Ars Pharmaceutica, 48(1), 5-17.",
    "COSADES. (s.f.). Historia y desarrollo institucional del CESFAM Villa Nonguén. "
    "Corporación de Salud y Desarrollo Social. https://www.cosades.cl",
    "Delgado, O., Anoz, L., Serrano, A., & Nicolás, J. (2007). Conciliación "
    "farmacoterapéutica como garantía de continuidad asistencial. Sociedad Española de "
    "Farmacia Hospitalaria.",
    "Diario Concepción. (2017, 11 de marzo). La historia del Cesfam que rompió "
    "paradigmas de la salud pública. https://www.diarioconcepcion.cl",
    "Espinosa García, J., Prados Torres, J. D., Leiva Fernández, F., & Barnestein "
    "Fonseca, P. (2023). Adherencia terapéutica de los pacientes con riesgo "
    "cardiovascular en Atención Primaria. Proyecto REAAP. Medicina de Familia. "
    "SEMERGEN, 49(6), 102016.",
    "Faúndez Navarrete, P. A. (2020). Evaluación del programa de Atención Farmacéutica "
    "del CESFAM Villa Nonguén [Memoria de título, Universidad de Concepción].",
    "Haynes, R. B., Taylor, D. W., & Sackett, D. L. (1979). Compliance in health care. "
    "Johns Hopkins University Press.",
    "Lambert, K., Bugnon, O., Del Giorno, R., Schneider, M. P., & PANDIA-IRIS Study "
    "Group. (2024). The differential impact of a 6- versus 12-month pharmacist-led "
    "interprofessional medication adherence program on medication adherence in patients "
    "with diabetic kidney disease: The randomized PANDIA-IRIS study. Frontiers in "
    "Pharmacology, 15, 1294436.",
    "Manzano García, M., & Morillo Verdugo, R. (2018). Aprendizaje y aplicación del "
    "modelo de Atención Farmacéutica CMO para residentes de farmacia hospitalaria. "
    "Sociedad Española de Farmacia Hospitalaria.",
    "Martín Alfonso, L. (2004). Acerca del concepto de adherencia terapéutica. Revista "
    "Cubana de Salud Pública, 30(4).",
    "Ministerio de Salud de Chile. (2017a). Encuesta Nacional de Salud 2016-2017. "
    "Gobierno de Chile.",
    "Ministerio de Salud de Chile. (2017b). Orientación técnica Programa de Salud "
    "Cardiovascular. Gobierno de Chile.",
    "Ministerio de Salud de Chile. (2018a). Guía de Atención Farmacéutica y seguimiento "
    "farmacoterapéutico en APS. Gobierno de Chile.",
    "Ministerio de Salud de Chile. (2018b). Orientación técnica Fondo de Farmacia. "
    "Gobierno de Chile.",
    "Ministerio de Salud de Chile. (2020). Resolución Exenta N.º 51: Aprueba Programa "
    "Fondo de Farmacia para Enfermedades Crónicas no Transmisibles. Gobierno de Chile.",
    "Morillo-Verdugo, R., Robustillo-Cortés, M. A., Navarro-Ruiz, A., "
    "Sánchez-Rubio Ferrández, J., & Fernández-Espínola, S. (2022). Clinical impact of "
    "the capacity-motivation-opportunity pharmacist-led intervention in people living "
    "with HIV in Spain, 2019-2020. Journal of Multidisciplinary Healthcare, 15, "
    "1203-1211.",
    "Morisky, D. E., Green, L. W., & Levine, D. M. (1986). Concurrent and predictive "
    "validity of a self-reported measure of medication adherence. Medical Care, 24(1), "
    "67-74.",
    "Sánchez-Yáñez, E., Manzano-García, M., & Morillo-Verdugo, R. (2023). Application "
    "of CMO (capacity, motivation, and opportunity) methodology in pharmaceutical care "
    "to optimize the pharmacotherapy in older people living with HIV. DISPIMDINAC "
    "project. Revista Española de Quimioterapia, 36(6), 584-591.",
]
for r in sorted(refs, key=lambda s: s.lower()):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(1); p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(cap_terms(r)); run.font.name = FONT; run.font.size = Pt(12)

# --------------------------------------------------------------- ANEXOS
heading("ANEXOS")
para("Los anexos reúnen los principales instrumentos de apoyo diseñados para el "
     "protocolo de Atención Farmacéutica domiciliaria del seminario de título. Son "
     "formatos de trabajo pensados para usarse durante la visita y el seguimiento en el "
     "CESFAM, y pueden ajustarse a los registros propios del establecimiento.")

heading("Anexo 1. Test de adherencia de Morisky-Green-Levine", level=2)
caption("Tabla", "Preguntas del test de Morisky-Green-Levine.")
make_table(["N.º", "Pregunta", "Respuesta esperada"], [
    ["1", "¿Olvida alguna vez tomar sus medicamentos?", "No"],
    ["2", "¿Toma los medicamentos a la hora indicada?", "Sí"],
    ["3", "Cuando se siente bien, ¿deja de tomarlos?", "No"],
    ["4", "Si alguna vez le sientan mal, ¿deja de tomarlos?", "No"],
], col_widths=[1.5, 9.5, 4.0], font_size=11)
para("Se considera adherente al paciente que responde de la forma esperada las cuatro "
     "preguntas. Cualquier otra combinación indica problemas de adherencia que conviene "
     "abordar en la entrevista.", size=11)

heading("Anexo 2. Clasificación de los RNM según el Consenso de Granada", level=2)
caption("Tabla", "Clasificación de los resultados negativos asociados a la medicación.")
make_table(["Dimensión", "Pregunta que orienta", "Ejemplo"], [
    ["Necesidad", "¿El paciente usa los medicamentos que necesita?",
     "Toma un fármaco sin indicación o falta uno que requiere."],
    ["Efectividad", "¿El tratamiento logra el objetivo buscado?",
     "No alcanza la meta terapéutica pese a usar el medicamento."],
    ["Seguridad", "¿El tratamiento es seguro para el paciente?",
     "Aparece una reacción adversa o una interacción relevante."],
], col_widths=[3.0, 6.0, 6.0], font_size=11)

heading("Anexo 3. Otros instrumentos del protocolo", level=2)
bullet("Ficha de conciliación farmacéutica: compara la medicación indicada con la que "
       "el paciente realmente usa y deja constancia de los cambios acordados.")
bullet("Ficha de seguimiento farmacoterapéutico: registra los antecedentes del "
       "paciente, los PRM y RNM detectados, y el resultado de las intervenciones.")
bullet("Hoja de esquema de horarios de medicación: ordena, en una tabla simple, qué "
       "medicamento tomar y a qué hora, como apoyo para el paciente y su cuidador.")
bullet("Protocolo de llamada telefónica: guion breve para contactar y agendar al "
       "paciente antes de la visita domiciliaria.")

doc.save("Informe de Internado - Farmacia Asistencial y APS - CESFAM Villa Nonguen.docx")
print("Documento generado correctamente.")
